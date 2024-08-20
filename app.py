import streamlit as st
import requests
from bs4 import BeautifulSoup
from fpdf import FPDF
from docx import Document
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

# Configure Streamlit page settings
st.set_page_config(
    page_title="Web Scraper",
    page_icon=":spider:",  # Favicon emoji
    layout="centered",  # Page layout option
)
# Custom CSS for styling
st.markdown("""
    <style>
    /* Background color */
    .main {
        background-color: #f5f5f5;
    }
    /* Title and headers */
    h1, h2, h3, h4 {
        color: #4B778D;
        font-family: 'Segoe UI', sans-serif;
    }
    /* Center the title */
    .css-10trblm {
        text-align: center;
    }
    /* Customize the sidebar */
    .css-1d391kg {
        background-color: #1F1F1F;
        color: #FFFFFF;
    }
    /* Customize the buttons */
    .stButton button {
        background-color: #4B778D;
        color: white;
        font-size: 18px;
        border-radius: 10px;
    }
    /* Image styling */
    img {
        border-radius: 10px;
    }
    </style>
""", unsafe_allow_html=True)

# Streamlit app title with a stylish header
st.markdown("<h1 style='text-align: center; color: #4B778D;'>🕸️ScrapeIt: Web Data Extractor</h1>", unsafe_allow_html=True)

# Sidebar for additional features
st.sidebar.title("App Options")
dark_mode = st.sidebar.checkbox("🌙 Dark Mode")

# Apply dark mode theme dynamically, including the sidebar
if dark_mode:
    st.markdown("""
    <style>
    .main {
        background-color: #1e1e1e;
        color: #f5f5f5;
    }
    .css-1d391kg {
        background-color: #333333;
        color: #f5f5f5;
    }
    h1, h2, h3, h4, p {
        color: #00ced1;
    }
    </style>
    """, unsafe_allow_html=True)

# Input: Website URL
url = st.text_input("Enter the URL of the website you want to scrape:")

# Input: Specify additional parameters if needed (like number of pages)
num_pages = st.number_input("Enter the number of pages to scrape (if applicable):", min_value=1, step=1, value=1)

# Input: Set the image width
image_width = st.slider("Set the image width (in pixels):", min_value=50, max_value=1000, value=300)

# Element Selection
st.sidebar.subheader("Select Elements to Scrape")
scrape_title = st.sidebar.checkbox("Scrape Title", value=True)
scrape_headers = st.sidebar.checkbox("Scrape Headers (h1, h2, h3, etc.)", value=True)
scrape_paragraphs = st.sidebar.checkbox("Scrape Paragraphs", value=True)
scrape_anchors = st.sidebar.checkbox("Scrape Anchor Tags (Links)", value=True)
scrape_images = st.sidebar.checkbox("Scrape Images", value=True)
scrape_comments = st.sidebar.checkbox("Scrape Comments", value=False)

if st.button("✨ Scrape Content ✨"):
    if url:
        # Function to scrape content from the webpage
        def scrape_content(url, num_pages):
            all_data = {
                'titles': [],
                'headers': [],
                'paragraphs': [],
                'anchors': [],
                'comments': [],
                'images': []
            }
            for page in range(1, num_pages + 1):
                paginated_url = f"{url}?page={page}"  # Adjust this for pagination if necessary
                response = requests.get(paginated_url)
                
                if response.status_code == 200:
                    soup = BeautifulSoup(response.content, 'html.parser')
                    
                    # Scrape the title
                    if scrape_title:
                        title = soup.title.get_text(strip=True) if soup.title else "No title found"
                        all_data['titles'].append(title)

                    # Scrape headers (h1, h2, h3, etc.)
                    if scrape_headers:
                        headers = [header.get_text(strip=True) for header in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6'])]
                        all_data['headers'].extend(headers)

                    # Scrape paragraphs (body content)
                    if scrape_paragraphs:
                        paragraphs = [p.get_text(strip=True) for p in soup.find_all('p')]
                        all_data['paragraphs'].extend(paragraphs)

                    # Scrape anchor tags (links)
                    if scrape_anchors:
                        anchors = [a.get('href') for a in soup.find_all('a') if a.get('href')]
                        all_data['anchors'].extend(anchors)

                    # Scrape comments (update selector according to the actual structure)
                    if scrape_comments:
                        comments = soup.find_all('div', class_='comment-class')  # Replace with actual tag and class for comments
                        for comment in comments:
                            all_data['comments'].append(comment.get_text(strip=True))

                    # Scrape images (src attribute)
                    if scrape_images:
                        images = soup.find_all('img')
                        for img in images:
                            img_src = img.get('src')
                            if img_src:
                                # Ensure the image URL is absolute
                                if not img_src.startswith('http'):
                                    img_src = requests.compat.urljoin(url, img_src)
                                all_data['images'].append(img_src)
                else:
                    st.error(f"Failed to retrieve page {page}. Status code: {response.status_code}")
                    break
            
            return all_data

        # Run the scraper and display the content
        with st.spinner("Scraping content..."):
            scraped_data = scrape_content(url, num_pages)
        
        if scraped_data:
            # Display the scraped content based on user selection
            if scrape_title and scraped_data['titles']:
                st.subheader("Page Titles")
                for idx, title in enumerate(scraped_data['titles'], 1):
                    st.markdown(f"<h3 style='color: #4B778D;'>{idx}. {title}</h3>", unsafe_allow_html=True)

            if scrape_headers and scraped_data['headers']:
                st.subheader("Headers")
                for header in scraped_data['headers']:
                    st.markdown(f"<p style='font-size: 18px;'>{header}</p>", unsafe_allow_html=True)

            if scrape_paragraphs and scraped_data['paragraphs']:
                st.subheader("Body Content (Paragraphs)")
                for paragraph in scraped_data['paragraphs']:
                    st.write(paragraph)

            if scrape_anchors and scraped_data['anchors']:
                st.subheader("Anchor Tags (Links)")
                for anchor in scraped_data['anchors']:
                    st.write(f"[{anchor}]({anchor})")

            if scrape_comments and scraped_data['comments']:
                st.subheader("Comments")
                for idx, comment in enumerate(scraped_data['comments'], 1):
                    st.markdown(f"<p style='background-color: #EFEFEF; padding: 10px; border-radius: 5px;'>{idx}. {comment}</p>", unsafe_allow_html=True)

            if scrape_images and scraped_data['images']:
                st.subheader("Images")
                for img_src in scraped_data['images']:
                    st.image(img_src, caption="Scraped Image", width=image_width)

            # Sidebar download options
            st.sidebar.subheader("Download Options")
            
            # Convert content to PDF
            def generate_pdf_with_reportlab(scraped_data):
                buffer = BytesIO()
                pdf = canvas.Canvas(buffer, pagesize=letter)
    
                # Title
                pdf.setFont("Helvetica-Bold", 16)
                pdf.drawString(100, 750, "Web Scraped Content")

                # Add Titles
                y_position = 700
                pdf.setFont("Helvetica-Bold", 12)
                for idx, title in enumerate(scraped_data['titles'], 1):
                    pdf.drawString(100, y_position, f"Title {idx}: {title}")
                    y_position -= 20
    
                # Add Headers
                pdf.setFont("Helvetica-Bold", 12)
                for header in scraped_data['headers']:
                    pdf.drawString(100, y_position, header)
                    y_position -= 20
    
                # Add Paragraphs
                pdf.setFont("Helvetica", 12)
                for paragraph in scraped_data['paragraphs']:
                    pdf.drawString(100, y_position, paragraph)
                    y_position -= 20
    
                pdf.save()
                buffer.seek(0)
                return buffer

            # Convert content to DOCX
            def generate_docx(scraped_data):
                doc = Document()
                doc.add_heading('Web Scraped Content', 0)
                
                # Add titles
                for idx, title in enumerate(scraped_data['titles'], 1):
                    doc.add_heading(f"Title {idx}: {title}", level=1)
                
                # Add headers
                for header in scraped_data['headers']:
                    doc.add_heading(header, level=2)
                
                # Add paragraphs
                for paragraph in scraped_data['paragraphs']:
                    doc.add_paragraph(paragraph)

                # Add anchor tags
                for anchor in scraped_data['anchors']:
                    doc.add_paragraph(anchor)

                # Add comments
                for idx, comment in enumerate(scraped_data['comments'], 1):
                    doc.add_paragraph(f"Comment {idx}: {comment}", style='Italic')
                
                # Convert DOCX to bytes
                doc_bytes = BytesIO()
                doc.save(doc_bytes)
                doc_bytes.seek(0)
                return doc_bytes

            # Download PDF
            pdf_file = generate_pdf_with_reportlab(scraped_data)
            st.sidebar.download_button(
            label=" 📄 Download PDF",
            data=pdf_file.getvalue(),  # Use getvalue() to extract bytes
            file_name="scraped_content.pdf",
            mime="application/pdf"
            )

            # Download DOCX
            docx_file = generate_docx(scraped_data)
            st.sidebar.download_button(
                label="📝  Download DOCX",
                data=docx_file.getvalue(),  # Use getvalue() to extract bytes
                file_name="scraped_content.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        else:
            st.warning("No content found or there was an error during scraping.")
    else:
        st.warning("Please enter a valid URL.")
